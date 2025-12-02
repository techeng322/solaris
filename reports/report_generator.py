"""
Report generator for insolation and KEO calculation results.
Generates formatted reports with diagrams, tables, and compliance information.
"""

from typing import List, Dict, Optional
from datetime import date
from pathlib import Path
import yaml

from models.calculation_result import BuildingCalculationResult, WindowCalculationResult
from models.building import Building


class ReportGenerator:
    """Generates formatted reports from calculation results."""
    
    def __init__(self, config_path: Optional[str] = None):
        """
        Initialize report generator.
        
        Args:
            config_path: Path to configuration file
        """
        self.config = self._load_config(config_path)
        self.output_format = self.config.get('reports', {}).get('format', 'pdf')
        self.include_diagrams = self.config.get('reports', {}).get('include_diagrams', True)
        self.include_stamps = self.config.get('reports', {}).get('include_stamps', True)
    
    def _load_config(self, config_path: Optional[str]) -> Dict:
        """Load configuration from file."""
        if config_path is None:
            config_path = 'config.yaml'
        
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                return yaml.safe_load(f) or {}
        except FileNotFoundError:
            return {}
    
    def generate_report(
        self,
        building_result: BuildingCalculationResult,
        output_path: str,
        building: Optional[Building] = None
    ) -> str:
        """
        Generate complete report for building calculation results.
        
        Args:
            building_result: Building calculation results
            output_path: Path to save report
            building: Building model (optional, for additional context)
        
        Returns:
            Path to generated report file
        """
        # Detect format from file extension (user's choice in file dialog)
        output_path_obj = Path(output_path)
        file_ext = output_path_obj.suffix.lower()
        
        # Determine format from extension, fallback to config if no extension
        if file_ext == '.pdf':
            format_type = 'pdf'
        elif file_ext == '.html' or file_ext == '.htm':
            format_type = 'html'
        elif file_ext == '.docx':
            format_type = 'docx'
        else:
            # No extension or unknown extension - use config default
            format_type = self.output_format
            # If still no format, default to PDF
            if not format_type:
                format_type = 'pdf'
        
        # Generate report in detected format
        if format_type == 'pdf':
            return self._generate_pdf_report(building_result, output_path, building)
        elif format_type == 'html':
            return self._generate_html_report(building_result, output_path, building)
        elif format_type == 'docx':
            return self._generate_docx_report(building_result, output_path, building)
        else:
            raise ValueError(f"Unsupported report format: {format_type}")
    
    def _generate_pdf_report(
        self,
        building_result: BuildingCalculationResult,
        output_path: str,
        building: Optional[Building]
    ) -> str:
        """Generate PDF report."""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.pdfgen import canvas
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        
        # Create PDF document
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=20*mm,
            leftMargin=20*mm,
            topMargin=20*mm,
            bottomMargin=20*mm
        )
        
        # Container for the 'Flowable' objects
        elements = []
        styles = getSampleStyleSheet()
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        title = Paragraph("Расчет инсоляции и КЕО", title_style)
        elements.append(title)
        elements.append(Spacer(1, 12))
        
        # Building information
        building_info = [
            ['Здание:', building_result.building_name],
            ['Дата расчета:', building_result.calculation_date.strftime('%d.%m.%Y') if building_result.calculation_date else 'N/A'],
        ]
        if building:
            building_info.append(['Количество окон:', str(building.get_total_windows())])
        
        info_table = Table(building_info, colWidths=[60*mm, 120*mm])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        elements.append(info_table)
        elements.append(Spacer(1, 20))
        
        # Summary
        summary = building_result.get_compliance_summary()
        summary_text = f"""
        <b>Сводка по зданию:</b><br/>
        Всего окон: {summary['total_windows']}<br/>
        Соответствующих требованиям: {summary['compliant_windows']}<br/>
        Не соответствующих требованиям: {summary['non_compliant_windows']}<br/>
        Процент соответствия: {summary['compliance_rate']*100:.1f}%
        """
        elements.append(Paragraph(summary_text, styles['Normal']))
        elements.append(Spacer(1, 20))
        
        # Window results
        elements.append(Paragraph("<b>Результаты по окнам:</b>", styles['Heading2']))
        elements.append(Spacer(1, 12))
        
        for window_result in building_result.window_results:
            # Window header
            window_name = window_result.window_name or window_result.window_id
            window_header = f"Окно: {window_name} (ID: {window_result.window_id})"
            elements.append(Paragraph(window_header, styles['Heading3']))
            
            # Insolation results
            if window_result.insolation_result:
                ins = window_result.insolation_result
                ins_text = f"""
                <b>Инсоляция:</b><br/>
                Продолжительность: {ins.duration_formatted}<br/>
                Соответствие требованиям: {'Да' if ins.meets_requirement else 'Нет'}
                """
                elements.append(Paragraph(ins_text, styles['Normal']))
            
            # KEO results
            if window_result.keo_result:
                keo = window_result.keo_result
                keo_text = f"""
                <b>КЕО:</b><br/>
                Общий КЕО: {keo.keo_total:.2f}%<br/>
                Соответствие требованиям: {'Да' if keo.meets_requirement else 'Нет'}
                """
                elements.append(Paragraph(keo_text, styles['Normal']))
            
            # Compliance status
            status_color = colors.green if window_result.is_compliant else colors.red
            status_text = f"<b>Статус:</b> {'Соответствует' if window_result.is_compliant else 'Не соответствует'}"
            status_para = Paragraph(status_text, styles['Normal'])
            elements.append(status_para)
            
            # Warnings
            if window_result.warnings:
                warnings_text = "<b>Предупреждения:</b><br/>" + "<br/>".join(window_result.warnings)
                elements.append(Paragraph(warnings_text, styles['Normal']))
            
            elements.append(Spacer(1, 20))
            elements.append(PageBreak())
        
        # Build PDF
        doc.build(elements)
        
        return output_path
    
    def _generate_html_report(
        self,
        building_result: BuildingCalculationResult,
        output_path: str,
        building: Optional[Building]
    ) -> str:
        """Generate HTML report."""
        html_content = f"""
        <!DOCTYPE html>
        <html lang="ru">
        <head>
            <meta charset="UTF-8">
            <title>Расчет инсоляции и КЕО</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                h1 {{ color: #1a1a1a; }}
                table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background-color: #f2f2f2; }}
                .compliant {{ color: green; }}
                .non-compliant {{ color: red; }}
            </style>
        </head>
        <body>
            <h1>Расчет инсоляции и КЕО</h1>
            <h2>Здание: {building_result.building_name}</h2>
            <p>Дата расчета: {building_result.calculation_date.strftime('%d.%m.%Y') if building_result.calculation_date else 'N/A'}</p>
        """
        
        # Add window results
        for window_result in building_result.window_results:
            window_name = window_result.window_name or window_result.window_id
            html_content += f"""
            <h3>Окно: {window_name}</h3>
            <table>
                <tr>
                    <th>Параметр</th>
                    <th>Значение</th>
                </tr>
            """
            
            if window_result.insolation_result:
                ins = window_result.insolation_result
                html_content += f"""
                <tr>
                    <td>Инсоляция</td>
                    <td>{ins.duration_formatted}</td>
                </tr>
                <tr>
                    <td>Соответствие инсоляции</td>
                    <td class="{'compliant' if ins.meets_requirement else 'non-compliant'}">
                        {'Да' if ins.meets_requirement else 'Нет'}
                    </td>
                </tr>
                """
            
            if window_result.keo_result:
                keo = window_result.keo_result
                html_content += f"""
                <tr>
                    <td>КЕО</td>
                    <td>{keo.keo_total:.2f}%</td>
                </tr>
                <tr>
                    <td>Соответствие КЕО</td>
                    <td class="{'compliant' if keo.meets_requirement else 'non-compliant'}">
                        {'Да' if keo.meets_requirement else 'Нет'}
                    </td>
                </tr>
                """
            
            html_content += f"""
                <tr>
                    <td>Общий статус</td>
                    <td class="{'compliant' if window_result.is_compliant else 'non-compliant'}">
                        {'Соответствует' if window_result.is_compliant else 'Не соответствует'}
                    </td>
                </tr>
            """
            
            html_content += "</table>"
        
        html_content += """
        </body>
        </html>
        """
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return output_path
    
    def _generate_docx_report(
        self,
        building_result: BuildingCalculationResult,
        output_path: str,
        building: Optional[Building]
    ) -> str:
        """Generate DOCX report."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
        except ImportError:
            # Fall back to HTML if python-docx is not available
            import logging
            logging.warning("python-docx not available, falling back to HTML")
            return self._generate_html_report(building_result, output_path.replace('.docx', '.html'), building)
        
        # Create document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
        
        # Title
        title = doc.add_heading('Расчет инсоляции и КЕО', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Building information
        doc.add_paragraph(f'Здание: {building_result.building_name}')
        if building_result.calculation_date:
            doc.add_paragraph(f'Дата расчета: {building_result.calculation_date.strftime("%d.%m.%Y")}')
        if building:
            doc.add_paragraph(f'Количество окон: {building.get_total_windows()}')
        
        doc.add_paragraph()  # Empty line
        
        # Summary
        summary = building_result.get_compliance_summary()
        summary_para = doc.add_paragraph()
        summary_para.add_run('Сводка по зданию:').bold = True
        doc.add_paragraph(f'Всего окон: {summary["total_windows"]}')
        doc.add_paragraph(f'Соответствующих требованиям: {summary["compliant_windows"]}')
        doc.add_paragraph(f'Не соответствующих требованиям: {summary["non_compliant_windows"]}')
        doc.add_paragraph(f'Процент соответствия: {summary["compliance_rate"]*100:.1f}%')
        
        doc.add_paragraph()  # Empty line
        
        # Window results
        doc.add_heading('Результаты по окнам', level=1)
        
        for window_result in building_result.window_results:
            window_name = window_result.window_name or window_result.window_id
            doc.add_heading(f'Окно: {window_name}', level=2)
            
            # Insolation results
            if window_result.insolation_result:
                ins = window_result.insolation_result
                doc.add_paragraph('Инсоляция:', style='List Bullet')
                doc.add_paragraph(f'  Продолжительность: {ins.duration_formatted}')
                doc.add_paragraph(f'  Соответствие требованиям: {"Да" if ins.meets_requirement else "Нет"}')
            
            # KEO results
            if window_result.keo_result:
                keo = window_result.keo_result
                doc.add_paragraph('КЕО:', style='List Bullet')
                doc.add_paragraph(f'  Общий КЕО: {keo.keo_total:.2f}%')
                doc.add_paragraph(f'  Соответствие требованиям: {"Да" if keo.meets_requirement else "Нет"}')
            
            # Compliance status
            status_para = doc.add_paragraph()
            status_run = status_para.add_run(f'Статус: {"Соответствует" if window_result.is_compliant else "Не соответствует"}')
            status_run.bold = True
            if window_result.is_compliant:
                status_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            else:
                status_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            
            # Warnings
            if window_result.warnings:
                doc.add_paragraph('Предупреждения:', style='List Bullet')
                for warning in window_result.warnings:
                    doc.add_paragraph(f'  - {warning}')
            
            doc.add_paragraph()  # Empty line between windows
            doc.add_page_break()
        
        # Save document
        doc.save(output_path)
        return output_path

